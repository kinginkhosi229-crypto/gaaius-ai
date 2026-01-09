from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Form, Depends, Request
from fastapi.responses import StreamingResponse, FileResponse, HTMLResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import asyncio
import base64
import io
import hashlib
import jwt
import requests
import json

# Load environment variables
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# API Keys
GROQ_API_KEY = os.environ.get('GROQ_API_KEY')
HF_TOKEN = os.environ.get('HF_TOKEN')
PAYPAL_CLIENT_ID = os.environ.get('PAYPAL_CLIENT_ID')
PAYPAL_SECRET = os.environ.get('PAYPAL_SECRET')
PAYFAST_MERCHANT_ID = os.environ.get('PAYFAST_MERCHANT_ID')
PAYFAST_MERCHANT_KEY = os.environ.get('PAYFAST_MERCHANT_KEY')
JWT_SECRET = os.environ.get('JWT_SECRET', 'default_secret')

# Initialize Groq client
from groq import Groq
groq_client = Groq(api_key=GROQ_API_KEY)

# Initialize Hugging Face client
from huggingface_hub import InferenceClient
hf_client = InferenceClient(api_key=HF_TOKEN)

# Import GAAIUS Build Brain v2
from gaaius_builder import (
    APP_TEMPLATES, 
    generate_blueprint, 
    quality_gate_v2, 
    GAAIUS_BUILD_PROMPT_V2,
    BLUEPRINT_SYSTEM_PROMPT,
    get_template_code,
    get_available_templates
)

# Create the main app
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer(auto_error=False)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============== MODELS ==============

class UserRegister(BaseModel):
    email: str
    password: str
    name: str = ""

class UserLogin(BaseModel):
    email: str
    password: str

class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    name: str = ""
    password_hash: str
    is_pro: bool = False
    pro_expires: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ChatMessage(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    role: str
    content: str
    model_used: Optional[str] = None
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChatRequest(BaseModel):
    session_id: str
    message: str
    
class ChatResponse(BaseModel):
    id: str
    content: str
    model_used: str
    timestamp: str

class ImageGenerationRequest(BaseModel):
    prompt: str
    session_id: Optional[str] = None

class ImageGenerationResponse(BaseModel):
    id: str
    prompt: str
    image_url: str
    model_used: str
    timestamp: str

class VideoGenerationRequest(BaseModel):
    prompt: str
    duration: int = 5
    style: str = "cinematic"
    session_id: Optional[str] = None

class VideoGenerationResponse(BaseModel):
    id: str
    prompt: str
    video_url: str
    model_used: str
    timestamp: str

class TTSRequest(BaseModel):
    text: str
    voice: str = "en"

class AudioGenerationRequest(BaseModel):
    prompt: str
    duration: int = 10
    type: str = "music"  # music, sfx, ambient

class FileGenerationRequest(BaseModel):
    prompt: str
    file_type: str  # code, document, data, config

class ProjectCreate(BaseModel):
    name: str
    description: str = ""
    type: str = "web"  # web, api, data

class Session(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str = "New Chat"
    user_id: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# ============== AUTH HELPERS ==============

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def create_token(user_id: str, email: str, is_pro: bool) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "is_pro": is_pro,
        "exp": datetime.now(timezone.utc) + timedelta(days=30)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    if not credentials:
        return None
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=["HS256"])
        user = await db.users.find_one({"id": payload["user_id"]}, {"_id": 0})
        return user
    except:
        return None

# ============== AUTH ROUTES ==============

@api_router.post("/auth/register")
async def register(data: UserRegister):
    existing = await db.users.find_one({"email": data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user = User(
        email=data.email,
        name=data.name,
        password_hash=hash_password(data.password)
    )
    await db.users.insert_one(user.model_dump())
    token = create_token(user.id, user.email, user.is_pro)
    
    return {"token": token, "user": {"id": user.id, "email": user.email, "name": user.name, "is_pro": user.is_pro}}

@api_router.post("/auth/login")
async def login(data: UserLogin):
    user = await db.users.find_one({"email": data.email}, {"_id": 0})
    if not user or user["password_hash"] != hash_password(data.password):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Check pro status
    is_pro = user.get("is_pro", False)
    if is_pro and user.get("pro_expires"):
        if datetime.fromisoformat(user["pro_expires"]) < datetime.now(timezone.utc):
            is_pro = False
            await db.users.update_one({"id": user["id"]}, {"$set": {"is_pro": False}})
    
    token = create_token(user["id"], user["email"], is_pro)
    return {"token": token, "user": {"id": user["id"], "email": user["email"], "name": user.get("name", ""), "is_pro": is_pro}}

@api_router.get("/auth/me")
async def get_me(user = Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return {"id": user["id"], "email": user["email"], "name": user.get("name", ""), "is_pro": user.get("is_pro", False)}

# ============== PAYMENT ROUTES ==============

@api_router.post("/payment/paypal/create")
async def create_paypal_order(user = Depends(get_current_user)):
    """Create PayPal order for Pro subscription"""
    try:
        # Get PayPal access token
        auth = base64.b64encode(f"{PAYPAL_CLIENT_ID}:{PAYPAL_SECRET}".encode()).decode()
        token_response = requests.post(
            "https://api-m.paypal.com/v1/oauth2/token",
            headers={"Authorization": f"Basic {auth}", "Content-Type": "application/x-www-form-urlencoded"},
            data="grant_type=client_credentials"
        )
        access_token = token_response.json().get("access_token")
        
        # Create order
        order_response = requests.post(
            "https://api-m.paypal.com/v2/checkout/orders",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json"
            },
            json={
                "intent": "CAPTURE",
                "purchase_units": [{
                    "amount": {"currency_code": "USD", "value": "1.00"},
                    "description": "GAAIUS AI Pro - 1 Month"
                }]
            }
        )
        return order_response.json()
    except Exception as e:
        logger.error(f"PayPal create error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/payment/paypal/capture/{order_id}")
async def capture_paypal_order(order_id: str, user = Depends(get_current_user)):
    """Capture PayPal payment and activate Pro"""
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        auth = base64.b64encode(f"{PAYPAL_CLIENT_ID}:{PAYPAL_SECRET}".encode()).decode()
        token_response = requests.post(
            "https://api-m.paypal.com/v1/oauth2/token",
            headers={"Authorization": f"Basic {auth}", "Content-Type": "application/x-www-form-urlencoded"},
            data="grant_type=client_credentials"
        )
        access_token = token_response.json().get("access_token")
        
        capture_response = requests.post(
            f"https://api-m.paypal.com/v2/checkout/orders/{order_id}/capture",
            headers={"Authorization": f"Bearer {access_token}", "Content-Type": "application/json"}
        )
        result = capture_response.json()
        
        if result.get("status") == "COMPLETED":
            # Activate Pro for 30 days
            expires = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
            await db.users.update_one(
                {"id": user["id"]},
                {"$set": {"is_pro": True, "pro_expires": expires}}
            )
            await db.payments.insert_one({
                "id": str(uuid.uuid4()),
                "user_id": user["id"],
                "provider": "paypal",
                "order_id": order_id,
                "amount": 1.00,
                "currency": "USD",
                "status": "completed",
                "timestamp": datetime.now(timezone.utc).isoformat()
            })
            return {"success": True, "message": "Pro activated!", "expires": expires}
        
        raise HTTPException(status_code=400, detail="Payment not completed")
    except Exception as e:
        logger.error(f"PayPal capture error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/payment/payfast/create")
async def create_payfast_payment(user = Depends(get_current_user)):
    """Generate PayFast payment URL"""
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    payment_id = str(uuid.uuid4())
    
    # PayFast payment data
    data = {
        "merchant_id": PAYFAST_MERCHANT_ID,
        "merchant_key": PAYFAST_MERCHANT_KEY,
        "return_url": f"https://gaaius-studio-1.preview.emergentagent.com/?payment=success&id={payment_id}",
        "cancel_url": "https://gaaius-studio-1.preview.emergentagent.com/?payment=cancelled",
        "notify_url": f"https://gaaius-studio-1.preview.emergentagent.com/api/payment/payfast/notify",
        "amount": "18.00",  # ~$1 in ZAR
        "item_name": "GAAIUS AI Pro - 1 Month",
        "custom_str1": user["id"],
        "custom_str2": payment_id
    }
    
    # Generate signature
    param_string = "&".join([f"{k}={v}" for k, v in sorted(data.items()) if k != "signature"])
    signature = hashlib.md5(param_string.encode()).hexdigest()
    data["signature"] = signature
    
    # Store pending payment
    await db.payments.insert_one({
        "id": payment_id,
        "user_id": user["id"],
        "provider": "payfast",
        "amount": 18.00,
        "currency": "ZAR",
        "status": "pending",
        "timestamp": datetime.now(timezone.utc).isoformat()
    })
    
    return {"payment_url": "https://www.payfast.co.za/eng/process", "data": data}

@api_router.post("/payment/payfast/notify")
async def payfast_notify(request: Request):
    """PayFast ITN callback"""
    try:
        form_data = await request.form()
        data = dict(form_data)
        
        if data.get("payment_status") == "COMPLETE":
            user_id = data.get("custom_str1")
            payment_id = data.get("custom_str2")
            
            expires = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
            await db.users.update_one({"id": user_id}, {"$set": {"is_pro": True, "pro_expires": expires}})
            await db.payments.update_one({"id": payment_id}, {"$set": {"status": "completed"}})
        
        return {"status": "ok"}
    except Exception as e:
        logger.error(f"PayFast notify error: {e}")
        return {"status": "error"}

@api_router.get("/payment/config")
async def get_payment_config():
    """Get payment configuration for frontend"""
    return {
        "paypal_client_id": PAYPAL_CLIENT_ID,
        "payfast_merchant_id": PAYFAST_MERCHANT_ID,
        "pro_price_usd": 1.00,
        "pro_price_zar": 18.00
    }

# ============== BASIC ROUTES ==============

@api_router.get("/")
async def root():
    return {"message": "GAAIUS AI Backend Running", "status": "operational"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "groq": bool(GROQ_API_KEY), "huggingface": bool(HF_TOKEN)}

# ============== SESSION ROUTES ==============

@api_router.post("/sessions", response_model=dict)
async def create_session(name: str = "New Chat", user = Depends(get_current_user)):
    session = Session(name=name, user_id=user["id"] if user else None)
    doc = session.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    doc['updated_at'] = doc['updated_at'].isoformat()
    await db.sessions.insert_one(doc)
    return {"id": session.id, "name": session.name, "created_at": doc['created_at']}

@api_router.put("/sessions/{session_id}")
async def update_session(session_id: str, data: dict):
    """Update session name"""
    update_data = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if "name" in data:
        update_data["name"] = data["name"]
    await db.sessions.update_one({"id": session_id}, {"$set": update_data})
    return {"status": "updated"}

@api_router.get("/sessions")
async def get_sessions(user = Depends(get_current_user)):
    query = {"user_id": user["id"]} if user else {}
    sessions = await db.sessions.find(query, {"_id": 0}).sort("updated_at", -1).to_list(100)
    return sessions

@api_router.delete("/sessions/{session_id}")
async def delete_session(session_id: str):
    await db.sessions.delete_one({"id": session_id})
    await db.messages.delete_many({"session_id": session_id})
    return {"status": "deleted"}

# ============== CHAT ROUTES ==============

@api_router.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest, user = Depends(get_current_user)):
    try:
        history = await db.messages.find({"session_id": request.session_id}, {"_id": 0}).sort("timestamp", 1).to_list(50)
        
        messages = [{"role": "system", "content": "You are GAAIUS AI, a powerful unified AI assistant. You can help with text conversations, image generation, video creation, audio synthesis, and file generation. Be helpful, creative, and engaging."}]
        
        for msg in history:
            messages.append({"role": msg['role'], "content": msg['content']})
        
        messages.append({"role": "user", "content": request.message})
        
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.7,
            max_tokens=2048
        )
        
        response_content = completion.choices[0].message.content
        model_used = "Groq Llama 3.3 70B"
        
        # Save messages
        user_msg = ChatMessage(session_id=request.session_id, role="user", content=request.message)
        user_doc = user_msg.model_dump()
        user_doc['timestamp'] = user_doc['timestamp'].isoformat()
        await db.messages.insert_one(user_doc)
        
        assistant_msg = ChatMessage(session_id=request.session_id, role="assistant", content=response_content, model_used=model_used)
        assistant_doc = assistant_msg.model_dump()
        assistant_doc['timestamp'] = assistant_doc['timestamp'].isoformat()
        await db.messages.insert_one(assistant_doc)
        
        await db.sessions.update_one({"id": request.session_id}, {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}})
        
        return ChatResponse(id=assistant_msg.id, content=response_content, model_used=model_used, timestamp=assistant_doc['timestamp'])
        
    except Exception as e:
        logger.error(f"Chat error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/chat/{session_id}/history")
async def get_chat_history(session_id: str):
    messages = await db.messages.find({"session_id": session_id}, {"_id": 0}).sort("timestamp", 1).to_list(1000)
    return messages

# ============== IMAGE GENERATION ==============

@api_router.post("/image/generate", response_model=ImageGenerationResponse)
async def generate_image(request: ImageGenerationRequest, user = Depends(get_current_user)):
    try:
        import requests as req
        from PIL import Image as PILImage
        import urllib.parse
        
        # Use Pollinations.ai - 100% FREE, no signup, no API key needed!
        encoded_prompt = urllib.parse.quote(request.prompt)
        API_URL = f"https://image.pollinations.ai/prompt/{encoded_prompt}?width=1024&height=1024&nologo=true"
        
        response = req.get(API_URL, timeout=120, allow_redirects=True)
        
        if response.status_code != 200 or 'image' not in response.headers.get('content-type', ''):
            raise Exception(f"Pollinations API error: {response.status_code}")
        
        image_bytes = response.content
        
        # Save image
        gen_id = str(uuid.uuid4())
        img_filename = f"{gen_id}.jpg"
        img_path = ROOT_DIR / "static" / img_filename
        (ROOT_DIR / "static").mkdir(exist_ok=True)
        
        # Convert bytes to image and save
        image = PILImage.open(io.BytesIO(image_bytes))
        image.save(img_path, format='JPEG', quality=90)
        
        image_url = f"/api/static/{img_filename}"
        model_used = "Pollinations AI (Free)"
        timestamp = datetime.now(timezone.utc).isoformat()
        
        await db.generations.insert_one({
            "id": gen_id, "type": "image", "prompt": request.prompt, "url": image_url,
            "model_used": model_used, "session_id": request.session_id, "timestamp": timestamp
        })
        
        return ImageGenerationResponse(id=gen_id, prompt=request.prompt, image_url=image_url, model_used=model_used, timestamp=timestamp)
        
    except Exception as e:
        logger.error(f"Image generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Image generation failed: {str(e)}")

# ============== VIDEO GENERATION ==============

from video_engine import VideoEngine, StoryVideoEngine

video_engine = VideoEngine(hf_token=HF_TOKEN, groq_api_key=GROQ_API_KEY, output_dir=ROOT_DIR / "static" / "videos")
story_video_engine = StoryVideoEngine(hf_token=HF_TOKEN, groq_api_key=GROQ_API_KEY, output_dir=ROOT_DIR / "static" / "videos")

@api_router.post("/video/generate", response_model=VideoGenerationResponse)
async def generate_video(request: VideoGenerationRequest, user = Depends(get_current_user)):
    try:
        result = await video_engine.generate_video(
            prompt=request.prompt, duration=min(request.duration, 30), fps=8, style=request.style
        )
        
        video_filename = Path(result["video_path"]).name
        video_url = f"/api/static/videos/{video_filename}"
        model_used = f"GAAIUS Video Engine ({request.style})"
        gen_id = result["video_id"]
        timestamp = datetime.now(timezone.utc).isoformat()
        
        await db.generations.insert_one({
            "id": gen_id, "type": "video", "prompt": request.prompt, "url": video_url,
            "model_used": model_used, "session_id": request.session_id, "timestamp": timestamp
        })
        
        return VideoGenerationResponse(id=gen_id, prompt=request.prompt, video_url=video_url, model_used=model_used, timestamp=timestamp)
        
    except Exception as e:
        logger.error(f"Video generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Video generation failed: {str(e)}")

@api_router.post("/video/generate-story")
async def generate_story_video(request: dict, user = Depends(get_current_user)):
    try:
        result = await story_video_engine.generate_story_video(
            story_prompt=request.get("prompt", ""),
            chapters=min(request.get("chapters", 3), 5),
            duration_per_chapter=min(request.get("duration_per_chapter", 8), 15),
            style=request.get("style", "cinematic")
        )
        
        video_filename = Path(result["video_path"]).name
        video_url = f"/api/static/videos/{video_filename}"
        
        return {"id": result["video_id"], "video_url": video_url, "chapters": result.get("chapters", []), "total_duration": result.get("total_duration", 0)}
    except Exception as e:
        logger.error(f"Story video error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ============== AUDIO GENERATION (HuggingFace TTS/STT) ==============

@api_router.post("/tts")
async def text_to_speech(request: TTSRequest, user = Depends(get_current_user)):
    """Text-to-Speech using HuggingFace"""
    try:
        # Try multiple TTS models as fallback
        models_to_try = [
            "espnet/kan-bayashi_ljspeech_vits",
            "facebook/mms-tts-eng",
            "microsoft/speecht5_tts"
        ]
        
        audio = None
        last_error = None
        
        for model in models_to_try:
            try:
                audio = hf_client.text_to_speech(request.text, model=model)
                if audio:
                    break
            except Exception as e:
                last_error = e
                continue
        
        if not audio:
            raise last_error or Exception("TTS failed with all models")
        
        return StreamingResponse(io.BytesIO(audio), media_type="audio/wav", headers={"Content-Disposition": "attachment; filename=speech.wav"})
        
    except Exception as e:
        logger.error(f"TTS error: {e}")
        raise HTTPException(status_code=500, detail=f"TTS failed: {str(e)}")

@api_router.post("/stt")
async def speech_to_text(audio: UploadFile = File(...), user = Depends(get_current_user)):
    """Speech-to-Text using HuggingFace Whisper"""
    try:
        audio_content = await audio.read()
        
        # Use Whisper via HuggingFace
        result = hf_client.automatic_speech_recognition(audio_content, model="openai/whisper-large-v3")
        
        text = result.get("text", "") if isinstance(result, dict) else str(result)
        return {"text": text, "model_used": "Whisper (HuggingFace)"}
        
    except Exception as e:
        logger.error(f"STT error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/audio/generate")
async def generate_audio(request: AudioGenerationRequest, user = Depends(get_current_user)):
    """Generate audio narration - AI creates stories, narrates text, or reads whatever you type"""
    try:
        from gtts import gTTS
        
        gen_id = str(uuid.uuid4())
        audio_filename = f"{gen_id}.mp3"
        audio_path = ROOT_DIR / "static" / "audio" / audio_filename
        (ROOT_DIR / "static" / "audio").mkdir(parents=True, exist_ok=True)
        
        # Detect language hints in prompt
        prompt_lower = request.prompt.lower()
        lang = 'en'  # Default English
        
        # Language detection based on keywords
        if any(word in prompt_lower for word in ['spanish', 'español', 'espanol']):
            lang = 'es'
        elif any(word in prompt_lower for word in ['french', 'français', 'francais']):
            lang = 'fr'
        elif any(word in prompt_lower for word in ['german', 'deutsch']):
            lang = 'de'
        elif any(word in prompt_lower for word in ['italian', 'italiano']):
            lang = 'it'
        elif any(word in prompt_lower for word in ['portuguese', 'português']):
            lang = 'pt'
        elif any(word in prompt_lower for word in ['chinese', '中文']):
            lang = 'zh-CN'
        elif any(word in prompt_lower for word in ['japanese', '日本語']):
            lang = 'ja'
        elif any(word in prompt_lower for word in ['korean', '한국어']):
            lang = 'ko'
        elif any(word in prompt_lower for word in ['russian', 'русский']):
            lang = 'ru'
        elif any(word in prompt_lower for word in ['arabic', 'عربي']):
            lang = 'ar'
        elif any(word in prompt_lower for word in ['hindi', 'हिंदी']):
            lang = 'hi'
        
        # Detect if user wants a story or creative content
        is_story_request = any(word in prompt_lower for word in [
            'story', 'tell me', 'create a', 'write a', 'make a', 'narrate a',
            'story about', 'tale', 'adventure', 'explain', 'describe'
        ])
        
        # Extract duration if specified (e.g., "2 minutes", "30 seconds")
        import re
        duration_match = re.search(r'(\d+)\s*(minute|min|second|sec)', prompt_lower)
        target_words = 150  # Default ~1 minute
        if duration_match:
            num = int(duration_match.group(1))
            unit = duration_match.group(2)
            if 'min' in unit:
                target_words = num * 150  # ~150 words per minute
            else:
                target_words = max(30, num * 2)  # ~2 words per second
        
        if is_story_request:
            # Use AI to create the story/content
            completion = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": f"""You are a master storyteller and narrator. Create engaging, vivid content based on the user's request.
- If they ask for a story, write an entertaining story with characters and plot
- If they ask for an explanation, provide a clear and engaging explanation
- If they ask for a description, paint a vivid picture with words
- Target approximately {target_words} words
- Make it suitable for audio narration (no visual elements, emojis, or formatting)
- Use natural, flowing language that sounds great when read aloud"""},
                    {"role": "user", "content": request.prompt}
                ],
                temperature=0.8,
                max_tokens=min(4000, target_words * 2)
            )
            narration_text = completion.choices[0].message.content
        else:
            # Direct narration - just read what they typed or enhance slightly
            if len(request.prompt) < 20:
                narration_text = request.prompt
            else:
                # Light enhancement for better narration
                completion = groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[
                        {"role": "system", "content": "You are a narrator. Take the user's text and narrate it naturally. Keep the original meaning but make it flow well for audio. Do not add extra content, just narrate what they provided."},
                        {"role": "user", "content": f"Narrate this: {request.prompt}"}
                    ],
                    temperature=0.3,
                    max_tokens=1000
                )
                narration_text = completion.choices[0].message.content
        
        # Convert to speech using gTTS
        tts = gTTS(text=narration_text, lang=lang, slow=False)
        tts.save(str(audio_path))
        
        audio_url = f"/api/static/audio/{audio_filename}"
        timestamp = datetime.now(timezone.utc).isoformat()
        
        lang_names = {'en': 'English', 'es': 'Spanish', 'fr': 'French', 'de': 'German', 
                     'it': 'Italian', 'pt': 'Portuguese', 'zh-CN': 'Chinese', 'ja': 'Japanese',
                     'ko': 'Korean', 'ru': 'Russian', 'ar': 'Arabic', 'hi': 'Hindi'}
        
        await db.generations.insert_one({
            "id": gen_id, "type": "audio", "prompt": request.prompt, "url": audio_url,
            "content": narration_text, "language": lang_names.get(lang, 'English'), "timestamp": timestamp
        })
        
        return {"id": gen_id, "audio_url": audio_url, "content": narration_text, "language": lang_names.get(lang, 'English'), "timestamp": timestamp}
        
    except Exception as e:
        logger.error(f"Audio generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Audio generation failed: {str(e)}")

# ============== FILE GENERATION ==============

@api_router.post("/file/generate")
async def generate_file(request: FileGenerationRequest, user = Depends(get_current_user)):
    """Generate code/documents including PDF, DOCX, XLSX"""
    try:
        system_prompts = {
            "code": "You are an expert programmer. Generate clean, well-documented code based on the user's request. Output only the code, no explanations.",
            "document": "You are a professional document writer. Generate well-structured content with clear sections and paragraphs. Use markdown formatting with # for headers.",
            "data": "You are a data expert. Generate sample data in the exact format requested - JSON, CSV, XML. Output only the data.",
            "config": "You are a DevOps expert. Generate configuration files. Output only the config, no explanations."
        }
        
        prompt_lower = request.prompt.lower()
        ext = "txt"
        is_binary = False
        
        # Detect file format
        if "pdf" in prompt_lower:
            ext = "pdf"
            is_binary = True
        elif "docx" in prompt_lower or "word" in prompt_lower:
            ext = "docx"
            is_binary = True
        elif "xlsx" in prompt_lower or "excel" in prompt_lower:
            ext = "xlsx"
            is_binary = True
        elif request.file_type == "code":
            if "python" in prompt_lower or ".py" in prompt_lower:
                ext = "py"
            elif "javascript" in prompt_lower or ".js" in prompt_lower:
                ext = "js"
            elif "typescript" in prompt_lower or ".ts" in prompt_lower:
                ext = "ts"
            elif "html" in prompt_lower:
                ext = "html"
            elif "css" in prompt_lower:
                ext = "css"
            else:
                ext = "py"
        elif request.file_type == "document":
            if "html" in prompt_lower:
                ext = "html"
            elif "txt" in prompt_lower:
                ext = "txt"
            else:
                ext = "md"
        elif request.file_type == "data":
            if "csv" in prompt_lower:
                ext = "csv"
            elif "xml" in prompt_lower:
                ext = "xml"
            else:
                ext = "json"
        elif request.file_type == "config":
            if "yaml" in prompt_lower or "yml" in prompt_lower:
                ext = "yaml"
            elif "toml" in prompt_lower:
                ext = "toml"
            else:
                ext = "json"
        
        # Generate content using Groq
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompts.get(request.file_type, system_prompts["document"])},
                {"role": "user", "content": request.prompt}
            ],
            temperature=0.3,
            max_tokens=4096
        )
        
        content = completion.choices[0].message.content
        
        # Clean up code blocks
        if "```" in content:
            import re
            code_match = re.search(r'```[\w]*\n?([\s\S]*?)```', content)
            if code_match:
                content = code_match.group(1).strip()
        
        gen_id = str(uuid.uuid4())
        file_filename = f"{gen_id}.{ext}"
        file_path = ROOT_DIR / "static" / "files" / file_filename
        (ROOT_DIR / "static" / "files").mkdir(parents=True, exist_ok=True)
        
        # Generate binary files (PDF, DOCX, XLSX)
        if ext == "pdf":
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            
            doc = SimpleDocTemplate(str(file_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            for line in content.split('\n'):
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
            
            doc.build(story)
            
        elif ext == "docx":
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)
            doc.save(str(file_path))
            
        elif ext == "xlsx":
            from openpyxl import Workbook
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Data"
            
            for i, line in enumerate(content.split('\n'), 1):
                if line.strip():
                    cells = line.split(',') if ',' in line else [line]
                    for j, cell in enumerate(cells, 1):
                        ws.cell(row=i, column=j, value=cell.strip())
            wb.save(str(file_path))
        else:
            with open(file_path, "w") as f:
                f.write(content)
        
        file_url = f"/api/static/files/{file_filename}"
        timestamp = datetime.now(timezone.utc).isoformat()
        
        await db.generations.insert_one({
            "id": gen_id, "type": "file", "prompt": request.prompt, "url": file_url,
            "file_type": ext, "content": content if not is_binary else f"[{ext.upper()} file]", "timestamp": timestamp
        })
        
        return {"id": gen_id, "file_url": file_url, "content": content if not is_binary else f"[{ext.upper()} file generated]", "file_type": ext, "timestamp": timestamp}
        
    except Exception as e:
        logger.error(f"File generation error: {e}")
        raise HTTPException(status_code=500, detail=f"File generation failed: {str(e)}")

# ============== DOCUMENT STUDIO ==============

@api_router.post("/document/generate")
async def generate_document(data: dict, user = Depends(get_current_user)):
    """GAAIUS AI Document Studio - Generate professional documents"""
    try:
        prompt = data.get("prompt", "")
        doc_type = data.get("document_type", "pdf")
        current_content = data.get("current_content", "")
        doc_name = data.get("document_name", "document")
        
        # Document type specific prompts
        doc_prompts = {
            "invoice": """You are a professional invoice generator. Create a detailed, professional invoice with:
- Company/Sender information (placeholder for user to fill)
- Client/Bill To information
- Invoice number and date
- Itemized list with descriptions, quantities, rates, amounts
- Subtotal, Tax (if applicable), Total
- Payment terms and bank details
- Professional formatting with clear sections""",
            
            "contract": """You are a legal document writer. Create a comprehensive contract/agreement with:
- Party information sections
- Detailed terms and conditions
- Scope of work/services
- Payment terms
- Duration and termination clauses
- Confidentiality clause
- Dispute resolution
- Signature blocks
Use professional legal language.""",
            
            "proposal": """You are a business proposal writer. Create a compelling business proposal with:
- Executive Summary
- Problem Statement
- Proposed Solution
- Methodology/Approach
- Timeline and Milestones
- Team/Qualifications
- Pricing/Investment
- Terms and Conditions
- Call to Action
Use persuasive, professional language.""",
            
            "resume": """You are a professional CV/resume writer. Create a modern, ATS-friendly resume with:
- Contact Information
- Professional Summary
- Skills section
- Work Experience (reverse chronological)
- Education
- Certifications/Awards
Use action verbs and quantifiable achievements.""",
            
            "report": """You are a professional report writer. Create a detailed report with:
- Executive Summary
- Introduction
- Methodology
- Findings/Results
- Analysis
- Conclusions
- Recommendations
- References
Use clear headings and professional formatting.""",
            
            "letter": """You are a professional letter writer. Create a well-formatted business letter with:
- Date
- Recipient information
- Subject line
- Salutation
- Body paragraphs
- Closing
- Signature block
Use appropriate formal tone.""",
            
            "xlsx": """You are a spreadsheet/data expert. Create structured data that works well in Excel:
- Use comma-separated values
- Include clear headers in first row
- Use proper data types (numbers, dates, text)
- Include calculations/formulas descriptions
- Organize data logically""",
            
            "default": """You are a professional document writer. Create well-structured content with:
- Clear headings using markdown (# ## ###)
- Organized sections
- Professional language
- Proper formatting"""
        }
        
        system_prompt = doc_prompts.get(doc_type, doc_prompts["default"])
        
        # If editing existing content
        user_prompt = prompt
        if current_content:
            user_prompt = f"Current document content:\n{current_content[:2000]}\n\nUser request: {prompt}\n\nModify the document according to the request."
        
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.4,
            max_tokens=6000
        )
        
        content = completion.choices[0].message.content
        
        # Clean markdown code blocks
        if "```" in content:
            import re
            code_match = re.search(r'```[\w]*\n?([\s\S]*?)```', content)
            if code_match:
                content = code_match.group(1).strip()
        
        gen_id = str(uuid.uuid4())
        
        # Determine file extension
        ext_map = {
            "pdf": "pdf", "docx": "docx", "xlsx": "xlsx",
            "invoice": "pdf", "contract": "pdf", "proposal": "pdf",
            "resume": "pdf", "report": "pdf", "letter": "pdf",
            "presentation": "md"
        }
        ext = ext_map.get(doc_type, "md")
        
        file_filename = f"{gen_id}.{ext}"
        file_path = ROOT_DIR / "static" / "files" / file_filename
        (ROOT_DIR / "static" / "files").mkdir(parents=True, exist_ok=True)
        
        # Generate file based on type
        if ext == "pdf":
            try:
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib import colors
                from reportlab.lib.units import inch
                
                doc = SimpleDocTemplate(str(file_path), pagesize=A4, 
                    leftMargin=0.75*inch, rightMargin=0.75*inch,
                    topMargin=0.75*inch, bottomMargin=0.75*inch)
                styles = getSampleStyleSheet()
                
                # Custom styles
                styles.add(ParagraphStyle(name='Title', parent=styles['Heading1'], fontSize=18, spaceAfter=20))
                styles.add(ParagraphStyle(name='Subtitle', parent=styles['Heading2'], fontSize=14, spaceAfter=12))
                
                story = []
                
                for line in content.split('\n'):
                    line = line.strip()
                    if not line:
                        story.append(Spacer(1, 6))
                    elif line.startswith('# '):
                        story.append(Paragraph(line[2:], styles['Title']))
                    elif line.startswith('## '):
                        story.append(Paragraph(line[3:], styles['Subtitle']))
                    elif line.startswith('### '):
                        story.append(Paragraph(line[4:], styles['Heading3']))
                    elif line.startswith('- ') or line.startswith('* '):
                        story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
                    elif line.startswith(tuple('0123456789')):
                        story.append(Paragraph(line, styles['Normal']))
                    else:
                        story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 4))
                
                doc.build(story)
            except Exception as pdf_err:
                logger.error(f"PDF generation error: {pdf_err}")
                # Fallback to text file
                ext = "md"
                file_filename = f"{gen_id}.md"
                file_path = ROOT_DIR / "static" / "files" / file_filename
                with open(file_path, "w") as f:
                    f.write(content)
                    
        elif ext == "docx":
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                
                doc = Document()
                for line in content.split('\n'):
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.strip():
                        doc.add_paragraph(line)
                doc.save(str(file_path))
            except Exception as docx_err:
                logger.error(f"DOCX generation error: {docx_err}")
                ext = "md"
                file_filename = f"{gen_id}.md"
                file_path = ROOT_DIR / "static" / "files" / file_filename
                with open(file_path, "w") as f:
                    f.write(content)
                    
        elif ext == "xlsx":
            try:
                from openpyxl import Workbook
                from openpyxl.styles import Font, Alignment
                
                wb = Workbook()
                ws = wb.active
                ws.title = "Data"
                
                for i, line in enumerate(content.split('\n'), 1):
                    if line.strip():
                        cells = line.split(',') if ',' in line else line.split('\t') if '\t' in line else [line]
                        for j, cell in enumerate(cells, 1):
                            ws.cell(row=i, column=j, value=cell.strip())
                            if i == 1:  # Header row
                                ws.cell(row=i, column=j).font = Font(bold=True)
                wb.save(str(file_path))
            except Exception as xlsx_err:
                logger.error(f"XLSX generation error: {xlsx_err}")
                ext = "csv"
                file_filename = f"{gen_id}.csv"
                file_path = ROOT_DIR / "static" / "files" / file_filename
                with open(file_path, "w") as f:
                    f.write(content)
        else:
            with open(file_path, "w") as f:
                f.write(content)
        
        file_url = f"/api/static/files/{file_filename}"
        timestamp = datetime.now(timezone.utc).isoformat()
        
        await db.generations.insert_one({
            "id": gen_id, "type": "document", "prompt": prompt, "url": file_url,
            "document_type": doc_type, "content": content[:1000], "timestamp": timestamp
        })
        
        return {
            "id": gen_id, 
            "file_url": file_url, 
            "filename": f"{doc_name}.{ext}",
            "content": content,
            "document_type": doc_type,
            "message": f"Your {doc_type.upper()} document has been created! You can preview it and download.",
            "timestamp": timestamp
        }
        
    except Exception as e:
        logger.error(f"Document generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")

@api_router.post("/document/generate-professional")
async def generate_professional_document(data: dict, user = Depends(get_current_user)):
    """Generate REAL professional documents - proper invoices, quotes, receipts as PDFs"""
    try:
        prompt = data.get("prompt", "")
        doc_type = data.get("document_type", "invoice")
        doc_name = data.get("document_name", "Document")
        output_format = data.get("output_format", "pdf")
        
        # Professional document prompts that generate structured data
        professional_prompts = {
            "invoice": """You are a professional invoice generator. Based on the user request, generate invoice data in this EXACT JSON format:
{
  "invoice_number": "INV-2024-001",
  "date": "2024-01-15",
  "due_date": "2024-02-15",
  "company": {"name": "Your Company", "address": "123 Business St", "city": "City, State 12345", "email": "billing@company.com", "phone": "(555) 123-4567"},
  "client": {"name": "Client Name", "address": "456 Client Ave", "city": "City, State 67890", "email": "client@email.com"},
  "items": [{"description": "Service Description", "quantity": 1, "rate": 100.00, "amount": 100.00}],
  "subtotal": 100.00,
  "tax_rate": 0,
  "tax": 0,
  "total": 100.00,
  "notes": "Payment due within 30 days.",
  "payment_info": "Bank: Example Bank, Account: 1234567890"
}
Extract real details from the user request. Output ONLY valid JSON.""",

            "quotation": """You are a professional quotation generator. Based on the user request, generate quote data in this EXACT JSON format:
{
  "quote_number": "QT-2024-001",
  "date": "2024-01-15",
  "valid_until": "2024-02-15",
  "company": {"name": "Your Company", "address": "123 Business St", "city": "City, State 12345", "email": "sales@company.com", "phone": "(555) 123-4567"},
  "client": {"name": "Client Name", "company": "Client Company", "address": "456 Client Ave", "city": "City, State 67890"},
  "items": [{"description": "Item/Service Description", "quantity": 1, "unit_price": 100.00, "total": 100.00}],
  "subtotal": 100.00,
  "discount": 0,
  "tax": 0,
  "total": 100.00,
  "terms": "Quote valid for 30 days. 50% deposit required.",
  "notes": "Thank you for your inquiry."
}
Extract real details from the user request. Output ONLY valid JSON.""",

            "receipt": """You are a professional receipt generator. Based on the user request, generate receipt data in this EXACT JSON format:
{
  "receipt_number": "RCP-2024-001",
  "date": "2024-01-15",
  "company": {"name": "Your Company", "address": "123 Business St", "city": "City, State 12345", "phone": "(555) 123-4567"},
  "customer": {"name": "Customer Name", "email": "customer@email.com"},
  "items": [{"description": "Item/Service", "quantity": 1, "price": 100.00, "total": 100.00}],
  "subtotal": 100.00,
  "tax": 0,
  "total": 100.00,
  "payment_method": "Credit Card",
  "payment_reference": "TXN-123456",
  "notes": "Thank you for your payment!"
}
Extract real details from the user request. Output ONLY valid JSON.""",

            "xlsx": """You are a spreadsheet data generator. Based on the user request, generate spreadsheet data as CSV format:
- First row must be headers
- Use commas to separate columns
- Each row on a new line
- Include calculations descriptions where applicable
Output ONLY the CSV data, no explanations."""
        }
        
        system_prompt = professional_prompts.get(doc_type, professional_prompts.get("invoice"))
        
        # For xlsx, use different approach
        if doc_type == "xlsx":
            completion = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=4000
            )
            csv_content = completion.choices[0].message.content.strip()
            if "```" in csv_content:
                import re
                match = re.search(r'```(?:csv)?\n?([\s\S]*?)```', csv_content)
                if match:
                    csv_content = match.group(1).strip()
            
            # Generate Excel file
            gen_id = str(uuid.uuid4())
            file_filename = f"{gen_id}.xlsx"
            file_path = ROOT_DIR / "static" / "files" / file_filename
            (ROOT_DIR / "static" / "files").mkdir(parents=True, exist_ok=True)
            
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
            
            wb = Workbook()
            ws = wb.active
            ws.title = doc_name[:30]
            
            # Parse CSV and write to Excel
            rows = csv_content.split('\n')
            for i, row in enumerate(rows, 1):
                cells = row.split(',')
                for j, cell in enumerate(cells, 1):
                    ws.cell(row=i, column=j, value=cell.strip())
                    if i == 1:  # Header row
                        ws.cell(row=i, column=j).font = Font(bold=True)
                        ws.cell(row=i, column=j).fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
            
            # Auto-adjust column widths
            for col in ws.columns:
                max_length = max(len(str(cell.value or "")) for cell in col)
                ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)
            
            wb.save(str(file_path))
            file_url = f"/static/files/{file_filename}"
            
            return {
                "id": gen_id,
                "file_url": file_url,
                "filename": f"{doc_name}.xlsx",
                "content": csv_content,
                "format": "xlsx",
                "message": f"Excel spreadsheet created! Download it below."
            }
        
        # For invoices, quotes, receipts - generate structured data then create PDF
        if doc_type in ["invoice", "quotation", "receipt"]:
            completion = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=2000
            )
            
            json_response = completion.choices[0].message.content.strip()
            
            # Clean JSON if needed
            if "```" in json_response:
                import re
                match = re.search(r'```(?:json)?\n?([\s\S]*?)```', json_response)
                if match:
                    json_response = match.group(1).strip()
            
            try:
                doc_data = json.loads(json_response)
            except:
                # Fallback to basic text document
                return await generate_document(data, user)
            
            # Generate professional PDF
            gen_id = str(uuid.uuid4())
            file_filename = f"{gen_id}.pdf"
            file_path = ROOT_DIR / "static" / "files" / file_filename
            (ROOT_DIR / "static" / "files").mkdir(parents=True, exist_ok=True)
            
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_RIGHT, TA_CENTER
            
            doc = SimpleDocTemplate(str(file_path), pagesize=A4,
                leftMargin=0.5*inch, rightMargin=0.5*inch,
                topMargin=0.5*inch, bottomMargin=0.5*inch)
            
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(name='RightAlign', parent=styles['Normal'], alignment=TA_RIGHT))
            styles.add(ParagraphStyle(name='Center', parent=styles['Normal'], alignment=TA_CENTER))
            styles.add(ParagraphStyle(name='DocTitle', parent=styles['Heading1'], fontSize=28, spaceAfter=5, textColor=colors.HexColor('#333333')))
            styles.add(ParagraphStyle(name='CompanyName', parent=styles['Heading2'], fontSize=18, textColor=colors.HexColor('#333333'), spaceAfter=3))
            styles.add(ParagraphStyle(name='SmallText', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#666666')))
            styles.add(ParagraphStyle(name='SectionHeader', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#666666'), spaceBefore=15, spaceAfter=5))
            
            story = []
            
            if doc_type == "invoice":
                # Professional Invoice like the example images
                # Header with company name and INVOICE title
                header_data = [
                    [Paragraph(f"<b>{doc_data.get('company', {}).get('name', 'Your Company')}</b>", styles['CompanyName']), 
                     Paragraph("<b>INVOICE</b>", ParagraphStyle('InvTitle', parent=styles['Normal'], fontSize=32, alignment=TA_RIGHT, textColor=colors.HexColor('#2E7D32')))]
                ]
                header_table = Table(header_data, colWidths=[300, 190])
                header_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(header_table)
                
                # Company details
                company = doc_data.get("company", {})
                story.append(Paragraph(company.get('address', '123 Business Street'), styles['SmallText']))
                story.append(Paragraph(company.get('city', 'City, State 12345'), styles['SmallText']))
                story.append(Paragraph(f"Phone: {company.get('phone', '(555) 123-4567')} | Email: {company.get('email', 'info@company.com')}", styles['SmallText']))
                story.append(Spacer(1, 15))
                
                # Green accent bar with invoice details
                invoice_bar = [
                    [Paragraph(f"<b>Invoice No.</b><br/>{doc_data.get('invoice_number', 'INV-001')}", ParagraphStyle('BarText', fontSize=9, textColor=colors.white)),
                     Paragraph(f"<b>Issue Date</b><br/>{doc_data.get('date', 'Jan 15, 2024')}", ParagraphStyle('BarText', fontSize=9, textColor=colors.white)),
                     Paragraph(f"<b>Due Date</b><br/>{doc_data.get('due_date', 'Feb 15, 2024')}", ParagraphStyle('BarText', fontSize=9, textColor=colors.white)),
                     Paragraph(f"<b>Total Due</b><br/>${doc_data.get('total', 0):,.2f}", ParagraphStyle('BarText', fontSize=9, textColor=colors.white, alignment=TA_RIGHT))]
                ]
                bar_table = Table(invoice_bar, colWidths=[120, 120, 120, 130])
                bar_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (2, 0), colors.HexColor('#2E7D32')),  # Green
                    ('BACKGROUND', (3, 0), (3, 0), colors.HexColor('#333333')),  # Dark gray for total
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('ALIGN', (3, 0), (3, 0), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('PADDING', (0, 0), (-1, -1), 10),
                ]))
                story.append(bar_table)
                story.append(Spacer(1, 20))
                
                # Bill To section
                client = doc_data.get("client", {})
                story.append(Paragraph("<b>BILL TO:</b>", ParagraphStyle('BillTo', fontSize=10, textColor=colors.HexColor('#666666'))))
                story.append(Paragraph(f"<b>{client.get('name', 'Client Name')}</b>", styles['Normal']))
                story.append(Paragraph(client.get('address', ''), styles['SmallText']))
                story.append(Paragraph(client.get('city', ''), styles['SmallText']))
                if client.get('email'):
                    story.append(Paragraph(client.get('email', ''), styles['SmallText']))
                story.append(Spacer(1, 20))
                
                # Items table with professional styling
                items_data = [[
                    Paragraph("<b>DESCRIPTION</b>", ParagraphStyle('TH', fontSize=9, textColor=colors.white)),
                    Paragraph("<b>QTY</b>", ParagraphStyle('TH', fontSize=9, textColor=colors.white, alignment=TA_CENTER)),
                    Paragraph("<b>UNIT PRICE</b>", ParagraphStyle('TH', fontSize=9, textColor=colors.white, alignment=TA_RIGHT)),
                    Paragraph("<b>AMOUNT</b>", ParagraphStyle('TH', fontSize=9, textColor=colors.white, alignment=TA_RIGHT))
                ]]
                for item in doc_data.get("items", []):
                    items_data.append([
                        item.get('description', ''),
                        str(item.get('quantity', 1)),
                        f"${item.get('rate', 0):,.2f}",
                        f"${item.get('amount', 0):,.2f}"
                    ])
                
                items_table = Table(items_data, colWidths=[250, 50, 90, 100])
                items_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E7D32')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                    ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e0e0e0')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f8f8')]),
                    ('PADDING', (0, 0), (-1, -1), 8),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(items_table)
                story.append(Spacer(1, 15))
                
                # Totals section - right aligned
                totals_data = [
                    ["Subtotal:", f"${doc_data.get('subtotal', 0):,.2f}"],
                    [f"Tax ({doc_data.get('tax_rate', 0)}%):", f"${doc_data.get('tax', 0):,.2f}"],
                    ["", ""],  # Empty row for spacing
                    ["TOTAL:", f"${doc_data.get('total', 0):,.2f}"]
                ]
                totals_table = Table(totals_data, colWidths=[370, 120])
                totals_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                    ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, -1), (-1, -1), 12),
                    ('TEXTCOLOR', (0, -1), (-1, -1), colors.HexColor('#2E7D32')),
                    ('LINEABOVE', (1, -1), (1, -1), 2, colors.HexColor('#2E7D32')),
                    ('PADDING', (0, 0), (-1, -1), 3),
                ]))
                story.append(totals_table)
                story.append(Spacer(1, 30))
                
                # Notes and payment info
                if doc_data.get('notes'):
                    story.append(Paragraph("<b>Notes:</b>", styles['SectionHeader']))
                    story.append(Paragraph(doc_data.get('notes'), styles['SmallText']))
                if doc_data.get('payment_info'):
                    story.append(Spacer(1, 10))
                    story.append(Paragraph("<b>Payment Information:</b>", styles['SectionHeader']))
                    story.append(Paragraph(doc_data.get('payment_info'), styles['SmallText']))
                
                # Footer
                story.append(Spacer(1, 40))
                story.append(Paragraph("Thank you for your business!", ParagraphStyle('Footer', fontSize=10, alignment=TA_CENTER, textColor=colors.HexColor('#666666'))))
                    
            elif doc_type == "quotation":
                # Professional Quote like Stripe example
                header_data = [
                    [Paragraph(f"<b>{doc_data.get('company', {}).get('name', 'Your Company')}</b>", styles['CompanyName']), 
                     Paragraph("<b>QUOTE</b>", ParagraphStyle('QuoteTitle', parent=styles['Normal'], fontSize=32, alignment=TA_RIGHT, textColor=colors.HexColor('#2E7D32')))]
                ]
                header_table = Table(header_data, colWidths=[300, 190])
                story.append(header_table)
                
                company = doc_data.get("company", {})
                story.append(Paragraph(company.get('address', ''), styles['SmallText']))
                story.append(Paragraph(f"{company.get('city', '')} | {company.get('phone', '')} | {company.get('email', '')}", styles['SmallText']))
                story.append(Spacer(1, 15))
                
                # Quote info bar
                quote_bar = [
                    [Paragraph(f"<b>Quote No.</b><br/>{doc_data.get('quote_number', 'QT-001')}", ParagraphStyle('BarText', fontSize=9, textColor=colors.white)),
                     Paragraph(f"<b>Issue Date</b><br/>{doc_data.get('date', '')}", ParagraphStyle('BarText', fontSize=9, textColor=colors.white)),
                     Paragraph(f"<b>Valid Until</b><br/>{doc_data.get('valid_until', '')}", ParagraphStyle('BarText', fontSize=9, textColor=colors.white)),
                     Paragraph(f"<b>Total</b><br/>${doc_data.get('total', 0):,.2f}", ParagraphStyle('BarText', fontSize=9, textColor=colors.white, alignment=TA_RIGHT))]
                ]
                bar_table = Table(quote_bar, colWidths=[120, 120, 120, 130])
                bar_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (2, 0), colors.HexColor('#2E7D32')),
                    ('BACKGROUND', (3, 0), (3, 0), colors.HexColor('#333333')),
                    ('PADDING', (0, 0), (-1, -1), 10),
                ]))
                story.append(bar_table)
                story.append(Spacer(1, 20))
                
                # Quote For
                client = doc_data.get("client", {})
                story.append(Paragraph("<b>QUOTE FOR:</b>", styles['SectionHeader']))
                story.append(Paragraph(f"<b>{client.get('name', '')}</b>", styles['Normal']))
                if client.get('company'):
                    story.append(Paragraph(client.get('company'), styles['SmallText']))
                story.append(Paragraph(client.get('address', ''), styles['SmallText']))
                story.append(Spacer(1, 20))
                
                # Items
                items_data = [[
                    Paragraph("<b>DESCRIPTION</b>", ParagraphStyle('TH', fontSize=9, textColor=colors.white)),
                    Paragraph("<b>QTY</b>", ParagraphStyle('TH', fontSize=9, textColor=colors.white, alignment=TA_CENTER)),
                    Paragraph("<b>UNIT PRICE</b>", ParagraphStyle('TH', fontSize=9, textColor=colors.white, alignment=TA_RIGHT)),
                    Paragraph("<b>TOTAL</b>", ParagraphStyle('TH', fontSize=9, textColor=colors.white, alignment=TA_RIGHT))
                ]]
                for item in doc_data.get("items", []):
                    items_data.append([
                        item.get('description', ''),
                        str(item.get('quantity', 1)),
                        f"${item.get('unit_price', 0):,.2f}",
                        f"${item.get('total', 0):,.2f}"
                    ])
                
                items_table = Table(items_data, colWidths=[250, 50, 90, 100])
                items_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E7D32')),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e0e0e0')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f8f8')]),
                    ('PADDING', (0, 0), (-1, -1), 8),
                    ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                    ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
                ]))
                story.append(items_table)
                story.append(Spacer(1, 15))
                
                # Totals
                totals_data = [
                    ["Subtotal:", f"${doc_data.get('subtotal', 0):,.2f}"],
                    ["Discount:", f"-${doc_data.get('discount', 0):,.2f}"],
                    ["TOTAL:", f"${doc_data.get('total', 0):,.2f}"]
                ]
                totals_table = Table(totals_data, colWidths=[370, 120])
                totals_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                    ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, -1), (-1, -1), 12),
                    ('TEXTCOLOR', (0, -1), (-1, -1), colors.HexColor('#2E7D32')),
                    ('LINEABOVE', (1, -1), (1, -1), 2, colors.HexColor('#2E7D32')),
                ]))
                story.append(totals_table)
                story.append(Spacer(1, 25))
                
                if doc_data.get('terms'):
                    story.append(Paragraph("<b>Terms & Conditions:</b>", styles['SectionHeader']))
                    story.append(Paragraph(doc_data.get('terms'), styles['SmallText']))
                story.append(Spacer(1, 10))
                
                company = doc_data.get("company", {})
                story.append(Paragraph(f"<b>{company.get('name', 'Company Name')}</b>", styles['CompanyName']))
                story.append(Paragraph(f"{company.get('address', '')} | {company.get('city', '')}", styles['Normal']))
                story.append(Paragraph(f"Email: {company.get('email', '')} | Phone: {company.get('phone', '')}", styles['Normal']))
                story.append(Spacer(1, 20))
                
                # Quote info
                quote_info = [
                    ["Quote #:", doc_data.get('quote_number', 'QT-001')],
                    ["Date:", doc_data.get('date', '')],
                    ["Valid Until:", doc_data.get('valid_until', '')]
                ]
                info_table = Table(quote_info, colWidths=[100, 200])
                info_table.setStyle(TableStyle([('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold')]))
                story.append(info_table)
                story.append(Spacer(1, 20))
                
                # Client
                client = doc_data.get("client", {})
                story.append(Paragraph("<b>Quote For:</b>", styles['Normal']))
                story.append(Paragraph(client.get('name', ''), styles['Normal']))
                story.append(Paragraph(client.get('company', ''), styles['Normal']))
                story.append(Paragraph(client.get('address', ''), styles['Normal']))
                story.append(Spacer(1, 20))
                
                # Items
                items_data = [["Description", "Qty", "Unit Price", "Total"]]
                for item in doc_data.get("items", []):
                    items_data.append([
                        item.get('description', ''),
                        str(item.get('quantity', 1)),
                        f"${item.get('unit_price', 0):,.2f}",
                        f"${item.get('total', 0):,.2f}"
                    ])
                
                items_table = Table(items_data, colWidths=[280, 50, 80, 80])
                items_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#27ae60')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
                ]))
                story.append(items_table)
                story.append(Spacer(1, 10))
                
                # Totals
                totals_data = [
                    ["Subtotal:", f"${doc_data.get('subtotal', 0):,.2f}"],
                    ["Discount:", f"-${doc_data.get('discount', 0):,.2f}"],
                    ["TOTAL:", f"${doc_data.get('total', 0):,.2f}"]
                ]
                totals_table = Table(totals_data, colWidths=[390, 100])
                totals_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                    ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                    ('LINEABOVE', (0, -1), (-1, -1), 2, colors.black),
                ]))
                story.append(totals_table)
                story.append(Spacer(1, 30))
                
                if doc_data.get('terms'):
                    story.append(Paragraph(f"<b>Terms:</b> {doc_data.get('terms')}", styles['Normal']))
                    
            elif doc_type == "receipt":
                # Professional Receipt like shop receipt example
                # Center-aligned header
                story.append(Paragraph("<b>RECEIPT</b>", ParagraphStyle('ReceiptTitle', fontSize=28, alignment=TA_CENTER, textColor=colors.HexColor('#333333'), spaceAfter=10)))
                
                company = doc_data.get("company", {})
                story.append(Paragraph(f"<b>{company.get('name', 'Shop Name')}</b>", ParagraphStyle('ShopName', fontSize=14, alignment=TA_CENTER)))
                story.append(Paragraph(company.get('address', ''), ParagraphStyle('CenterSmall', fontSize=9, alignment=TA_CENTER, textColor=colors.HexColor('#666666'))))
                story.append(Paragraph(f"{company.get('city', '')} | {company.get('phone', '')}", ParagraphStyle('CenterSmall', fontSize=9, alignment=TA_CENTER, textColor=colors.HexColor('#666666'))))
                story.append(Spacer(1, 10))
                
                # Dashed line separator
                story.append(Paragraph("-" * 70, ParagraphStyle('Dash', fontSize=8, alignment=TA_CENTER, textColor=colors.HexColor('#cccccc'))))
                story.append(Spacer(1, 5))
                
                # Receipt info - centered
                story.append(Paragraph(f"<b>Receipt No:</b> {doc_data.get('receipt_number', 'RCP-001')}", ParagraphStyle('CenterNormal', fontSize=10, alignment=TA_CENTER)))
                story.append(Paragraph(f"Date: {doc_data.get('date', '')} | Time: {doc_data.get('time', '')}", ParagraphStyle('CenterSmall', fontSize=9, alignment=TA_CENTER, textColor=colors.HexColor('#666666'))))
                story.append(Spacer(1, 5))
                story.append(Paragraph("-" * 70, ParagraphStyle('Dash', fontSize=8, alignment=TA_CENTER, textColor=colors.HexColor('#cccccc'))))
                story.append(Spacer(1, 10))
                
                # Customer info
                customer = doc_data.get("customer", {})
                if customer.get('name'):
                    story.append(Paragraph(f"Customer: {customer.get('name', '')}", styles['SmallText']))
                    story.append(Spacer(1, 10))
                
                # Items table - receipt style
                items_data = [[
                    Paragraph("<b>Item</b>", ParagraphStyle('TH', fontSize=9)),
                    Paragraph("<b>Qty</b>", ParagraphStyle('TH', fontSize=9, alignment=TA_CENTER)),
                    Paragraph("<b>Price</b>", ParagraphStyle('TH', fontSize=9, alignment=TA_RIGHT)),
                    Paragraph("<b>Total</b>", ParagraphStyle('TH', fontSize=9, alignment=TA_RIGHT))
                ]]
                for item in doc_data.get("items", []):
                    items_data.append([
                        item.get('description', ''),
                        str(item.get('quantity', 1)),
                        f"${item.get('price', 0):,.2f}",
                        f"${item.get('total', 0):,.2f}"
                    ])
                
                items_table = Table(items_data, colWidths=[200, 50, 70, 80])
                items_table.setStyle(TableStyle([
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                    ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
                    ('LINEBELOW', (0, 0), (-1, 0), 1, colors.HexColor('#333333')),
                    ('LINEBELOW', (0, -1), (-1, -1), 0.5, colors.HexColor('#cccccc')),
                    ('PADDING', (0, 0), (-1, -1), 5),
                ]))
                story.append(items_table)
                story.append(Spacer(1, 10))
                
                # Separator
                story.append(Paragraph("-" * 70, ParagraphStyle('Dash', fontSize=8, alignment=TA_CENTER, textColor=colors.HexColor('#cccccc'))))
                
                # Totals
                totals_data = [
                    ["Subtotal:", f"${doc_data.get('subtotal', 0):,.2f}"],
                    ["Tax:", f"${doc_data.get('tax', 0):,.2f}"],
                ]
                totals_table = Table(totals_data, colWidths=[320, 80])
                totals_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                ]))
                story.append(totals_table)
                
                # Grand total - prominent
                story.append(Spacer(1, 5))
                total_data = [["TOTAL:", f"${doc_data.get('total', 0):,.2f}"]]
                total_table = Table(total_data, colWidths=[320, 80])
                total_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 14),
                    ('LINEABOVE', (0, 0), (-1, -1), 2, colors.HexColor('#333333')),
                    ('PADDING', (0, 0), (-1, -1), 5),
                ]))
                story.append(total_table)
                story.append(Spacer(1, 10))
                
                # Payment info
                story.append(Paragraph("-" * 70, ParagraphStyle('Dash', fontSize=8, alignment=TA_CENTER, textColor=colors.HexColor('#cccccc'))))
                story.append(Paragraph(f"Payment: {doc_data.get('payment_method', 'Cash')}", ParagraphStyle('CenterNormal', fontSize=10, alignment=TA_CENTER)))
                if doc_data.get('payment_reference'):
                    story.append(Paragraph(f"Ref: {doc_data.get('payment_reference', '')}", ParagraphStyle('CenterSmall', fontSize=9, alignment=TA_CENTER, textColor=colors.HexColor('#666666'))))
                story.append(Spacer(1, 15))
                
                # Thank you message
                if doc_data.get('notes'):
                    story.append(Paragraph(doc_data.get('notes'), ParagraphStyle('ThankYou', fontSize=10, alignment=TA_CENTER, textColor=colors.HexColor('#666666'))))
                else:
                    story.append(Paragraph("Thank you for your purchase!", ParagraphStyle('ThankYou', fontSize=10, alignment=TA_CENTER, textColor=colors.HexColor('#666666'))))
            
            doc.build(story)
            file_url = f"/static/files/{file_filename}"
            
            # Convert doc_data back to readable text for preview
            preview_content = json.dumps(doc_data, indent=2)
            
            return {
                "id": gen_id,
                "file_url": file_url,
                "filename": f"{doc_name}.pdf",
                "content": preview_content,
                "format": "pdf",
                "message": f"Professional {doc_type} created! Click DOWNLOAD to get your PDF."
            }
        
        # For other document types, use the regular generator
        return await generate_document(data, user)
        
    except Exception as e:
        logger.error(f"Professional document generation error: {e}")
        # Fallback to regular document generation
        return await generate_document(data, user)

@api_router.post("/document/download")
async def download_document_as_format(data: dict, user = Depends(get_current_user)):
    """Download document content in specified format"""
    try:
        content = data.get("content", "")
        doc_type = data.get("document_type", "document")
        doc_name = data.get("document_name", "Document")
        format = data.get("format", "pdf")
        
        gen_id = str(uuid.uuid4())
        file_path = ROOT_DIR / "static" / "files" / f"{gen_id}.{format}"
        (ROOT_DIR / "static" / "files").mkdir(parents=True, exist_ok=True)
        
        if format == "pdf":
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            
            doc = SimpleDocTemplate(str(file_path), pagesize=A4,
                leftMargin=0.75*inch, rightMargin=0.75*inch,
                topMargin=0.75*inch, bottomMargin=0.75*inch)
            styles = getSampleStyleSheet()
            story = []
            
            for line in content.split('\n'):
                if line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
            
            doc.build(story)
            
        elif format == "docx":
            from docx import Document
            doc = Document()
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            doc.save(str(file_path))
            
        elif format == "xlsx":
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            for i, line in enumerate(content.split('\n'), 1):
                if line.strip():
                    cells = line.split(',') if ',' in line else [line]
                    for j, cell in enumerate(cells, 1):
                        ws.cell(row=i, column=j, value=cell.strip())
            wb.save(str(file_path))
            
        else:
            with open(file_path, 'w') as f:
                f.write(content)
        
        from fastapi.responses import FileResponse
        return FileResponse(
            path=str(file_path),
            filename=f"{doc_name}.{format}",
            media_type='application/octet-stream'
        )
        
    except Exception as e:
        logger.error(f"Document download error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ============== PROJECTS ==============

@api_router.post("/projects")
async def create_project(data: ProjectCreate, user = Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Login required")
    
    project = {
        "id": str(uuid.uuid4()),
        "user_id": user["id"],
        "name": data.name,
        "description": data.description,
        "type": data.type,
        "files": {},
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    await db.projects.insert_one(project)
    # Return without _id
    project.pop("_id", None)
    return project

@api_router.get("/projects")
async def get_projects(user = Depends(get_current_user)):
    if not user:
        return []
    projects = await db.projects.find({"user_id": user["id"]}, {"_id": 0}).sort("updated_at", -1).to_list(100)
    return projects

@api_router.get("/projects/{project_id}")
async def get_project(project_id: str, user = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project

@api_router.put("/projects/{project_id}/files")
async def update_project_files(project_id: str, files: dict, user = Depends(get_current_user)):
    await db.projects.update_one(
        {"id": project_id},
        {"$set": {"files": files, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "updated"}

# ============== GAAIUS BUILD BRAIN (Production-Grade AI Builder) ==============

# GAAIUS BUILD BRAIN - The Intelligence Layer
GAAIUS_SYSTEM_PROMPT = """SYSTEM: GAAIUS AI BUILDER - PRODUCTION-GRADE APPLICATION BUILDER

You are NOT a demo generator.
You are NOT creating minimal examples.
You are a PRODUCTION-GRADE application builder.

MANDATORY RULES:
1. NEVER generate minimal, ugly, or demo-quality UI
2. ALWAYS prioritize UI/UX quality - this must look investor-ready
3. ALWAYS use modern UI patterns with proper spacing, typography, and icons
4. ALWAYS use Tailwind CSS with professional component patterns
5. The output must look like a real, launched product
6. If a request is vague, make STRONG design decisions - never simplify

DESIGN STANDARDS (NON-NEGOTIABLE):
- Spacing: Minimum 16px padding, proper margins
- Typography: Clear hierarchy (headings, subheadings, body)
- Colors: Professional color palette with proper contrast
- Icons: Use inline SVG icons or icon CDN
- Layout: Use CSS Grid/Flexbox for proper alignment
- Responsive: Must work on mobile, tablet, desktop
- Interactivity: Hover states, transitions, animations
- Images: Use real placeholder images (picsum.photos, placehold.co)

FORCED FRONTEND STACK:
- Tailwind CSS (via CDN)
- Modern JavaScript (ES6+)
- Google Fonts (Inter, Poppins, or similar)
- Lucide Icons or Heroicons (via CDN)
- Smooth transitions and animations

OUTPUT FORMAT:
Return ONLY a complete, production-ready HTML file.
NO markdown. NO explanations. NO code blocks.
Just pure HTML that runs perfectly."""

def compile_user_prompt(raw_prompt):
    """PROMPT COMPILER - Converts vague user input into detailed spec"""
    
    prompt_lower = raw_prompt.lower()
    
    # Detect app type and expand
    app_expansions = {
        "youtube": "modern video streaming platform with: dark theme, sidebar navigation (Home, Trending, Subscriptions, Library), top navbar with search bar and user avatar, main content area with video grid (thumbnails, titles, channel names, view counts, timestamps), video player page, comments section, related videos sidebar, channel pages, responsive design",
        "spotify": "music streaming application with: dark theme (#121212), left sidebar (Home, Search, Library, Playlists), top bar with navigation and user menu, main content area with album grids, artist pages, playlist views, bottom player bar (album art, play/pause/skip controls, progress bar, volume), search with filters, responsive design",
        "netflix": "video streaming service with: dark theme, top navbar (logo, navigation, search, notifications, profile), hero banner with featured content, horizontal scrolling content rows (Trending, Continue Watching, New Releases), hover preview cards with details, categories, responsive design",
        "twitter": "social media platform with: light/dark theme toggle, left sidebar (Home, Explore, Notifications, Messages, Profile), main feed with tweets (avatar, username, content, images, engagement buttons), right sidebar (search, trends, who to follow), compose tweet modal, responsive design",
        "instagram": "photo sharing app with: top navbar (logo, search, icons), stories row, main feed with posts (user info, image, like/comment/share buttons, caption), right sidebar with suggestions, profile pages with grid layout, bottom mobile navigation, responsive design",
        "amazon": "e-commerce platform with: top navbar (logo, search bar, account, cart), category navigation, hero carousel, product grids with cards (image, title, rating, price, Prime badge), filters sidebar, product detail pages, cart functionality, responsive design",
        "dashboard": "admin dashboard with: sidebar navigation (collapsible), top header (search, notifications, profile), main content with stat cards, charts (line, bar, pie), data tables with sorting/filtering, activity feed, responsive grid layout",
        "landing": "modern landing page with: navbar (logo, links, CTA button), hero section (headline, subheadline, CTA, image), features section (icon cards), testimonials, pricing table, FAQ accordion, footer, smooth scroll animations",
        "portfolio": "portfolio website with: navbar (name, links), hero with introduction, about section, skills/tech stack, projects gallery with hover effects, testimonials, contact form, social links, responsive design",
        "blog": "blog platform with: navbar (logo, categories, search), featured post hero, post grid with cards (image, category, title, excerpt, date), sidebar (about, categories, tags, newsletter), single post view, comments, responsive design",
        "chat": "messaging application with: sidebar with conversation list, main chat area with message bubbles, message input with emoji/attachment support, user presence indicators, search conversations, responsive design"
    }
    
    # Find matching expansion
    expansion = ""
    for key, value in app_expansions.items():
        if key in prompt_lower:
            expansion = value
            break
    
    # Build compiled prompt
    if expansion:
        compiled = f"""BUILD REQUEST: Create a complete, production-ready {raw_prompt}

DESIGN SPECIFICATION:
{expansion}

QUALITY REQUIREMENTS:
- UI must look investor-ready and professionally designed
- All sections must be fully implemented with realistic content
- Include working interactivity (hover states, click handlers, modals)
- Use high-quality placeholder images from picsum.photos
- Implement smooth animations and transitions
- Ensure full responsiveness for all screen sizes

OUTPUT: Complete HTML file only. No explanations."""
    else:
        compiled = f"""BUILD REQUEST: {raw_prompt}

DESIGN SPECIFICATION:
Create a modern, professional, production-ready implementation with:
- Clean, spacious layout with proper visual hierarchy
- Professional typography and color scheme
- Navigation, header, main content, and footer sections
- Interactive elements with hover states and transitions
- Realistic sample content and images
- Full mobile responsiveness

QUALITY REQUIREMENTS:
- Must look like a real, launched product
- Investor-ready quality
- No minimal demos or wireframes

OUTPUT: Complete HTML file only. No explanations."""
    
    return compiled

def quality_check(html_code):
    """QUALITY GATE - Checks if generated code meets standards"""
    issues = []
    score = 100
    
    # Check for Tailwind CSS
    if "tailwindcss" not in html_code.lower():
        issues.append("Missing Tailwind CSS")
        score -= 20
    
    # Check for proper structure
    if "<nav" not in html_code.lower() and "navbar" not in html_code.lower():
        issues.append("Missing navigation")
        score -= 10
    
    # Check for responsive classes
    responsive_classes = ["md:", "lg:", "sm:", "xl:"]
    has_responsive = any(rc in html_code for rc in responsive_classes)
    if not has_responsive:
        issues.append("Missing responsive design")
        score -= 15
    
    # Check for proper spacing
    spacing_classes = ["p-", "px-", "py-", "m-", "mx-", "my-", "gap-", "space-"]
    has_spacing = any(sc in html_code for sc in spacing_classes)
    if not has_spacing:
        issues.append("Missing proper spacing")
        score -= 10
    
    # Check for icons
    if "svg" not in html_code.lower() and "lucide" not in html_code.lower() and "heroicon" not in html_code.lower():
        issues.append("Missing icons")
        score -= 5
    
    # Check for images
    if "img" not in html_code.lower() and "background-image" not in html_code.lower():
        issues.append("Missing images")
        score -= 5
    
    # Check for interactivity
    if "onclick" not in html_code.lower() and "hover:" not in html_code and "transition" not in html_code.lower():
        issues.append("Missing interactivity")
        score -= 10
    
    # Check for proper font
    if "font-" not in html_code or "googleapis.com/css" not in html_code:
        issues.append("Missing custom fonts")
        score -= 5
    
    return {
        "score": max(0, score),
        "passed": score >= 70,
        "issues": issues
    }

@api_router.post("/build/generate")
async def build_generate(data: dict, user = Depends(get_current_user)):
    """GAAIUS BUILD BRAIN v2.0 - Blueprint-First Platform Assembler"""
    try:
        raw_prompt = data.get("prompt", "")
        current_code = data.get("current_code", "")
        template_key = data.get("template", None)
        use_blueprint = data.get("use_blueprint", True)
        
        # STEP 1: Generate Blueprint
        blueprint = generate_blueprint(raw_prompt, template_key)
        logger.info(f"Blueprint generated: {blueprint.get('template_name', 'Custom')}")
        
        # STEP 2: Try template-based generation first
        template_code = None
        if blueprint.get('template_used') and use_blueprint:
            template_code = get_template_code(
                blueprint['template_used'],
                blueprint.get('app_name', 'MyApp')
            )
        
        # STEP 3: Compile user prompt into detailed spec
        compiled_prompt = compile_user_prompt(raw_prompt)
        
        # STEP 4: Build with enhanced system prompt
        messages = [
            {"role": "system", "content": GAAIUS_BUILD_PROMPT_V2},
        ]
        
        # Include blueprint context
        blueprint_context = f"""
APP BLUEPRINT:
- Name: {blueprint.get('app_name', 'MyApp')}
- Type: {blueprint.get('app_type', 'custom')}
- Template: {blueprint.get('template_name', 'Custom')}
- Pages: {', '.join([p['name'] for p in blueprint.get('pages', [])])}
- Features: {', '.join(blueprint.get('features', []))}
- Theme: {blueprint.get('theme', 'dark-modern')}
"""
        
        if current_code and len(current_code) > 100:
            messages.append({"role": "user", "content": f"{blueprint_context}\n\nCURRENT CODE:\n{current_code}\n\nUPDATE REQUEST:\n{compiled_prompt}\n\nReturn the complete updated HTML file."})
        else:
            messages.append({"role": "user", "content": f"{blueprint_context}\n\nBUILD REQUEST:\n{compiled_prompt}"})
        
        # STEP 5: Generate code
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.4,
            max_tokens=8000
        )
        
        code = completion.choices[0].message.content
        
        # Clean up code blocks
        if "```" in code:
            code_match = re.search(r'```(?:html)?\n?([\s\S]*?)```', code)
            if code_match:
                code = code_match.group(1)
        
        code = code.strip()
        
        # Ensure it starts with DOCTYPE
        if not code.lower().startswith('<!doctype'):
            html_start = code.lower().find('<!doctype')
            if html_start == -1:
                html_start = code.lower().find('<html')
            if html_start > 0:
                code = code[html_start:]
        
        # STEP 6: Enhanced Quality Gate
        quality = quality_gate_v2(code, blueprint)
        
        # STEP 7: Auto-regenerate if quality too low
        if not quality["passed"] and quality["score"] < 60:
            logger.info(f"Quality score {quality['score']} too low, regenerating with stricter requirements...")
            
            issues_text = ', '.join([i['msg'] for i in quality.get('issues', [])])
            regenerate_prompt = f"""{compiled_prompt}

CRITICAL QUALITY REQUIREMENTS (Previous generation scored {quality['score']}/100):
Issues to fix: {issues_text}

MANDATORY FIXES:
1. Include Tailwind CSS CDN
2. Add proper semantic HTML structure (nav, header, main, aside, footer)
3. Use responsive classes (sm:, md:, lg:, xl:)
4. Add Lucide icons (include script src="https://unpkg.com/lucide@latest/dist/umd/lucide.js")
5. Use proper spacing (p-4, p-6, gap-4, etc.)
6. Add hover states and smooth transitions
7. Generate at least 3000 characters of code
8. Include real content, not placeholders

Generate a COMPLETE, PRODUCTION-READY implementation following the blueprint."""

            messages[-1] = {"role": "user", "content": regenerate_prompt}
            
            completion = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=messages,
                temperature=0.3,
                max_tokens=8000
            )
            
            code = completion.choices[0].message.content
            if "```" in code:
                code_match = re.search(r'```(?:html)?\n?([\s\S]*?)```', code)
                if code_match:
                    code = code_match.group(1)
            code = code.strip()
            
            quality = quality_gate_v2(code, blueprint)
        
        return {
            "code": code,
            "model_used": "Groq Llama 3.3 (GAAIUS BUILD BRAIN v2)",
            "quality_score": quality["score"],
            "quality_passed": quality["passed"],
            "quality_checks": quality.get("checks_passed", []),
            "quality_issues": quality.get("issues", []),
            "blueprint": {
                "app_name": blueprint.get("app_name"),
                "template": blueprint.get("template_name"),
                "app_type": blueprint.get("app_type")
            }
        }
        
    except Exception as e:
        logger.error(f"Build generate error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/build/generate-full")
async def build_generate_full(data: dict, user = Depends(get_current_user)):
    """Generate a full web project with multiple files"""
    try:
        prompt = data.get("prompt", "")
        current_files = data.get("current_files", {})
        project_type = data.get("project_type", "web")
        
        system_prompt = """You are an expert full-stack web developer. You build REAL, functional websites and applications.

When the user asks you to build something, you must:
1. Create complete, working HTML files with embedded Tailwind CSS
2. Create proper JavaScript for interactivity
3. Create CSS for custom styling
4. Make it fully functional - not demos or mockups

Output format: Return a JSON object with:
- "files": an object where keys are filenames and values are the complete file contents
- "message": a brief description of what you built

Example response format:
{
  "files": {
    "index.html": "<!DOCTYPE html>...",
    "script.js": "// JavaScript code...",
    "style.css": "/* CSS styles */"
  },
  "message": "I built a responsive landing page with..."
}

IMPORTANT:
- Use Tailwind CSS via CDN in HTML
- Make the code production-ready
- Include proper meta tags and structure
- Add real functionality, not placeholder text
- Output ONLY valid JSON, no markdown or explanations"""
        
        # Build context from current files
        files_context = ""
        if current_files:
            files_context = "Current project files:\n"
            for filename, content in current_files.items():
                files_context += f"\n--- {filename} ---\n{content[:500]}...\n"
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"{files_context}\n\nUser request: {prompt}\n\nGenerate the updated/new files as JSON."}
        ]
        
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.3,
            max_tokens=8000
        )
        
        response_text = completion.choices[0].message.content.strip()
        
        # Try to parse JSON response
        try:
            # Remove markdown code blocks if present
            if response_text.startswith("```"):
                response_text = response_text.split("```")[1]
                if response_text.startswith("json"):
                    response_text = response_text[4:]
            
            result = json.loads(response_text)
            return {
                "files": result.get("files", {}),
                "message": result.get("message", "Code updated!")
            }
        except json.JSONDecodeError:
            # If not valid JSON, treat as single HTML file update
            return {
                "files": {"index.html": response_text},
                "message": "I've updated your index.html"
            }
        
    except Exception as e:
        logger.error(f"Build generate-full error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ============== STATIC FILES ==============

@api_router.get("/static/{filename}")
async def serve_static(filename: str):
    file_path = ROOT_DIR / "static" / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path)

@api_router.get("/static/videos/{filename}")
async def serve_video(filename: str):
    file_path = ROOT_DIR / "static" / "videos" / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Video not found")
    return FileResponse(file_path, media_type="video/mp4")

@api_router.get("/static/audio/{filename}")
async def serve_audio(filename: str):
    file_path = ROOT_DIR / "static" / "audio" / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Audio not found")
    return FileResponse(file_path, media_type="audio/wav")

@api_router.get("/static/files/{filename}")
async def serve_file(filename: str):
    file_path = ROOT_DIR / "static" / "files" / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path)

# ============== GENERATIONS ==============

@api_router.get("/generations")
async def get_generations(gen_type: Optional[str] = None, limit: int = 20):
    query = {}
    if gen_type:
        query["type"] = gen_type
    generations = await db.generations.find(query, {"_id": 0}).sort("timestamp", -1).to_list(limit)
    return generations

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
